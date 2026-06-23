from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/ONX_ATLAS_V5_COMPLETE_EXECUTION_PACKAGE_v1.0.docx")


TRAINS = [
    {
        "id": "L",
        "name": "Analytics, Reporting & Visualization Layer",
        "objective": "Build comprehensive dashboards, charts, metrics, and reporting surfaces across the Dream-to-Evolution lifecycle.",
        "routes": "/dashboard, /analytics, /reports",
        "dependencies": "Train K lifecycle data, Train I data table and stat-card primitives, Train H layout shell",
        "wps": [
            ("Dashboard Overview", "Implement executive KPI cards for dreams, potentials, goals, active execution, and completion velocity."),
            ("Dream Conversion Funnel", "Visualize Dream to Potential to Goal to Milestone to Completion conversion with stage-level rates."),
            ("Goal Progress Charts", "Add line, bar, and circular progress charts for goal completion, milestone burnup, and trend comparison."),
            ("Execution Velocity & Trends", "Track tasks completed per interval with trend indicators, variance notes, and operator filters."),
            ("Report Generation & Export", "Generate reviewable PDF and CSV reports with selectable periods, metrics, and evidence references."),
        ],
    },
    {
        "id": "M",
        "name": "Outcome, Flourishing & Recognition Workspace",
        "objective": "Implement outcome capture, lessons learned, achievement recognition, and flourishing signals after goals complete.",
        "routes": "/outcomes, /flourishing, /recognition",
        "dependencies": "Train K goals and execution records, Train L analytics baseline",
        "wps": [
            ("Outcome Recording", "Capture completed goal outcomes with result summaries, evidence links, and stakeholder impact."),
            ("Lessons Learned", "Record lessons, decisions, and improvement notes linked to completed execution work."),
            ("Achievement Badges", "Issue recognition badges for milestones, streaks, completion quality, and contribution patterns."),
            ("Flourishing Timeline", "Render a personal and institutional growth timeline across outcomes and recognitions."),
            ("Outcome Review Dashboard", "Summarize outcomes, lessons, recognitions, and completion quality for leadership review."),
        ],
    },
    {
        "id": "N",
        "name": "Evolution, Review & Continuous Improvement Layer",
        "objective": "Convert completed outcomes into structured improvement loops, retrospectives, and next-cycle recommendations.",
        "routes": "/evolution, /retrospectives",
        "dependencies": "Train M outcomes and lessons, Train L reports",
        "wps": [
            ("Retrospective Workspace", "Create guided retrospectives for completed goals and trains with owner, finding, and action fields."),
            ("Improvement Backlog", "Maintain a backlog of improvement items with priority, evidence, assignment, and status."),
            ("Pattern Detection", "Surface recurring lessons, blockers, and strengths from prior outcomes using rule-based summaries."),
            ("Recommendation Review", "Convert lessons into proposed next actions with approval, rejection, and deferral states."),
            ("Evolution Closure Report", "Generate closure reports that bind outcomes, lessons, improvements, and next-cycle actions."),
        ],
    },
    {
        "id": "O",
        "name": "Knowledge Corpus Expansion & Editorial Governance",
        "objective": "Expand the knowledge corpus while adding editorial controls for quality, provenance, publication, and review.",
        "routes": "/knowledge, /admin/knowledge",
        "dependencies": "Existing corpus ingestion, admin roles, evidence standard",
        "wps": [
            ("Corpus Expansion Workflow", "Add submission, draft, review, approved, and published states for knowledge articles."),
            ("Editorial Review Queue", "Build an admin queue for review assignment, comments, rejection reasons, and approval history."),
            ("Provenance Metadata", "Capture source, owner, date, confidence, category, and evidence references for each article."),
            ("Publication Controls", "Support scheduled publishing, unpublishing, archival, and public visibility checks."),
            ("Corpus Quality Audit", "Report missing metadata, stale articles, duplicate topics, and evidence gaps."),
        ],
    },
    {
        "id": "P",
        "name": "Titan Persona Governance & Conversational Operations",
        "objective": "Move Titan persona operations from static definitions toward governed records, auditability, and safe response operations.",
        "routes": "/admin/titans, /titan-conclave, /ask",
        "dependencies": "Titan database schema, existing hard-coded personas, Gate 6 conversation layer",
        "wps": [
            ("Persona Registry Admin", "Build governed admin editing for Titan domain, tone, prompt, traits, and activation state."),
            ("Persona Versioning", "Persist persona versions with change author, reason, approval status, and rollback metadata."),
            ("Conversation Audit Trail", "Expose searchable conversation records with persona, source, prompt, response, and timestamp."),
            ("Safety Guardrails", "Add refusal, escalation, and confidence rules for unsafe or out-of-domain Titan responses."),
            ("Titan Operations Report", "Summarize usage, retrieval quality, escalations, and persona changes for Kimi review."),
        ],
    },
    {
        "id": "Q",
        "name": "Institution Administration & Role Operations",
        "objective": "Harden institution management, role assignment, membership workflows, and admin evidence controls.",
        "routes": "/admin, /institution, /settings",
        "dependencies": "Better Auth, RBAC helpers, admin UI components",
        "wps": [
            ("Member Lifecycle", "Implement invite, activate, suspend, reactivate, and remove flows for institution members."),
            ("Role Assignment Controls", "Add founder/admin/operator role management with permission checks and audit events."),
            ("Access Review", "Create periodic role review reports with stale access, high-privilege accounts, and exceptions."),
            ("Admin Activity Log", "Record administrative actions with actor, subject, timestamp, route, and result."),
            ("Institution Settings", "Provide governed institution profile, policy, notification, and operational setting controls."),
        ],
    },
    {
        "id": "R",
        "name": "Security Hardening & Compliance Evidence",
        "objective": "Strengthen authentication, authorization, session controls, audit evidence, and compliance readiness.",
        "routes": "/login, /settings/security, /admin/security",
        "dependencies": "Better Auth, RBAC, route protection, evidence standard",
        "wps": [
            ("Session Security", "Add session inventory, revocation, expiration visibility, and secure logout coverage."),
            ("Password Policy", "Enforce password strength, reset safeguards, rate limits, and failure evidence."),
            ("Authorization Tests", "Expand pass and fail authorization tests across protected routes and procedures."),
            ("Security Evidence Pack", "Produce security evidence files for auth flows, role checks, and access denial."),
            ("Compliance Review Dashboard", "Expose compliance status, open findings, critical gaps, and closure evidence."),
        ],
    },
    {
        "id": "S",
        "name": "Runtime Reliability, Health & Observability",
        "objective": "Add production-grade health, readiness, metrics, logging, and incident visibility.",
        "routes": "/api/health, /api/health/ready, /api/metrics, /admin/operations",
        "dependencies": "Existing health endpoints, deployment scripts, monitoring docs",
        "wps": [
            ("Health Endpoint Expansion", "Validate app, database, auth, tRPC, and corpus readiness with clear status payloads."),
            ("Metrics Surface", "Emit operational metrics for request health, latency, failures, and domain workflows."),
            ("Structured Logging", "Standardize logs with event type, severity, request context, actor, and correlation id."),
            ("Incident Workflow", "Connect alerts, escalation, incident notes, owner assignment, and closure evidence."),
            ("Reliability Report", "Generate uptime, latency, error, smoke test, and rollback readiness reports."),
        ],
    },
    {
        "id": "T",
        "name": "Staging, Deployment & Release Governance",
        "objective": "Formalize staging validation, deployment evidence, rollback readiness, and release approvals.",
        "routes": "/admin/releases, deployment scripts",
        "dependencies": "Build pipeline, staging smoke tests, evidence folders",
        "wps": [
            ("Release Checklist", "Create a governed checklist for build, lint, smoke, evidence, approval, and rollout readiness."),
            ("Staging Smoke Automation", "Expand smoke tests for public, protected, admin, API, health, and auth flows."),
            ("Rollback Validation", "Document and test rollback steps with owner, timing, command, and evidence output."),
            ("Deployment Evidence Builder", "Collect deployment logs, URLs, health checks, commit hashes, and approval notes."),
            ("Release Certification", "Generate final readiness certification for Founder and Kimi approval."),
        ],
    },
    {
        "id": "U",
        "name": "Data Integrity, Migration & Backup Assurance",
        "objective": "Protect application data through integrity checks, migration discipline, backup validation, and recovery evidence.",
        "routes": "/admin/data, database scripts",
        "dependencies": "Drizzle schema, Neon database, migration scripts",
        "wps": [
            ("Schema Integrity Check", "Verify table prefixes, required indexes, relations, and schema drift evidence."),
            ("Migration Discipline", "Add migration generation, review, execution, and rollback documentation gates."),
            ("Backup Verification", "Define backup presence, restore rehearsal, retention, and recovery-point evidence."),
            ("Data Quality Reports", "Find orphaned records, invalid states, missing owners, and inconsistent lifecycle links."),
            ("Recovery Runbook", "Create database recovery procedures with roles, timing, commands, and validation outputs."),
        ],
    },
    {
        "id": "V",
        "name": "Performance, Load & Scalability Readiness",
        "objective": "Demonstrate acceptable performance under realistic load and document scale constraints.",
        "routes": "all critical user and API routes",
        "dependencies": "Runtime metrics, staging deployment, analytics routes",
        "wps": [
            ("Performance Budget", "Define budgets for page load, API latency, DB query duration, and bundle growth."),
            ("Load Test Harness", "Run repeatable concurrent-user tests for public, authenticated, admin, and API journeys."),
            ("Database Query Review", "Identify slow queries, missing indexes, N+1 risks, and pagination defects."),
            ("Frontend Performance Audit", "Review bundle cost, render stability, layout shifts, and responsive performance."),
            ("Scalability Report", "Compile load results, bottlenecks, remediations, residual risks, and sign-off evidence."),
        ],
    },
    {
        "id": "W",
        "name": "User Acceptance, Training & Operational Enablement",
        "objective": "Prepare users and operators with acceptance scripts, training materials, support workflows, and adoption evidence.",
        "routes": "/help, /support, training docs",
        "dependencies": "Completed feature trains, evidence standard, support ownership docs",
        "wps": [
            ("UAT Script Library", "Create acceptance scripts for core user, admin, Titan, knowledge, and reporting workflows."),
            ("Operator Training Guide", "Produce training guidance for daily operation, evidence collection, and issue handling."),
            ("Support Workflow", "Define intake, triage, escalation, resolution, and closure for support requests."),
            ("Adoption Dashboard", "Track training completion, UAT status, open blockers, and user feedback."),
            ("Enablement Certification", "Issue readiness certification for users, operators, and administrators."),
        ],
    },
    {
        "id": "X",
        "name": "Founder Review, Kimi Verification & Final Acceptance",
        "objective": "Package the final review flow so Kimi and Founder can verify completion without ambiguity.",
        "routes": "/admin/acceptance, evidence folders",
        "dependencies": "All prior train evidence, D02/D04/D05/D06 authority workflow",
        "wps": [
            ("Acceptance Matrix Finalization", "Create a complete matrix for train, WP, criterion, evidence, reviewer, and result."),
            ("Kimi Verification Workspace", "Provide review queue, scoring, comments, exceptions, and sign-off capture."),
            ("Founder Review Packet", "Compile executive summary, evidence index, unresolved risks, and decision prompts."),
            ("Exception Management", "Track rejected items, remediation owners, due dates, retest evidence, and closure status."),
            ("Final Acceptance Order", "Generate D06-style final order with PASS/FAIL status and authorized next steps."),
        ],
    },
    {
        "id": "Y",
        "name": "Production Launch, Monitoring & Closure",
        "objective": "Execute production launch readiness, monitor launch behavior, close Atlas V5, and preserve evidence.",
        "routes": "production launch surfaces and monitoring endpoints",
        "dependencies": "Founder acceptance, release certification, runtime monitoring",
        "wps": [
            ("Production Launch Checklist", "Confirm domains, env vars, migrations, health, rollback, access, and approvals."),
            ("Launch Communications", "Prepare launch notice, stakeholder update, support window, and escalation contacts."),
            ("Post-Launch Monitoring", "Monitor health, latency, errors, auth, data writes, and user journey success."),
            ("Closure Evidence Archive", "Commit complete evidence index, retention notes, release hashes, and approvals."),
            ("Atlas V5 Closure", "Issue Atlas V5 completion status, final metrics, owner sign-off, and post-launch watch plan."),
        ],
    },
    {
        "id": "Z",
        "name": "Atlas V5 Completion Assurance & Handover",
        "objective": "Close the remaining 15-train sequence with final assurance, operating handover, and long-term stewardship controls.",
        "routes": "/admin/closure, /admin/handover, repository docs",
        "dependencies": "Trains L through Y, production launch monitoring, evidence archive",
        "wps": [
            ("Completion Integrity Audit", "Verify all 75 WPs, 750 criteria, evidence records, and closure documents are present."),
            ("Operational Handover", "Create role-specific handover notes for Founder, Kimi, operators, and future agents."),
            ("Stewardship Runbook", "Document post-Atlas care routines, review cadence, maintenance windows, and authority boundaries."),
            ("Archive Retention Package", "Package retained evidence, final documents, commit hashes, and audit indexes for 90-day retention."),
            ("V5 to V6 Transition Brief", "Prepare the next-horizon brief with open risks, deferred work, and candidate V6 trains."),
        ],
    },
]


CRITERION_TEMPLATES = [
    "Implementation is present in the repository with code, route, component, script, or document changes directly traceable to this WP.",
    "The user-facing or operator-facing workflow can be exercised end to end without mocked success states.",
    "RBAC and protected-route behavior are enforced where the WP touches authenticated, admin, founder, or operator surfaces.",
    "Input validation, empty state, loading state, and failure state behavior are implemented where relevant.",
    "Evidence is collected under /evidence with the correct EP/date/category naming convention and a git commit hash reference.",
    "Automated test or scripted verification exists and records PASS/FAIL output for the WP.",
    "Build verification passes with bun run build, and lint verification passes with bun run lint before submission.",
    "The WP avoids reopening or weakening earlier accepted trains, routes, schemas, or evidence commitments.",
    "Kimi can independently review the result using the linked evidence, route, command output, and acceptance notes.",
    "Founder-facing closure language states the final PASS/FAIL status, residual risks, and next authorized step.",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ONX ATLAS V5 COMPLETE EXECUTION PACKAGE")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Trains L through Z | 15 Trains | 75 Work Packages | 750 Acceptance Criteria")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    add_table(
        doc,
        ["Document", "Version", "Date", "Status", "Authority"],
        [[OUT.name, "v1.0", "2026-06-23", "ACTIVE", "Founder Directive / Kimi Review"]],
    )


def add_overview(doc: Document) -> None:
    doc.add_heading("Part 1: Atlas V5 Completion Overview", level=1)
    doc.add_paragraph(
        "This package finalizes the remaining Atlas V5 execution sequence as a single reviewable deliverable. "
        "The directive requests 15 trains, 75 work packages, and 750 binary acceptance criteria. Because L through Y "
        "contains only 14 letters, this finalized package completes the arithmetic by issuing Trains L through Z."
    )
    add_table(
        doc,
        ["Scope Item", "Count", "Acceptance Rule"],
        [
            ["Trains", "15", "Each train passes only at 50/50 criteria."],
            ["Work Packages", "75", "Each WP passes only at 10/10 criteria."],
            ["Acceptance Criteria", "750", "Binary PASS/FAIL. No partial credit."],
            ["Evidence", "Required for every WP", "Stored under /evidence/<EP-ID>/<YYYY-MM-DD>/<category>/."],
        ],
    )
    add_table(
        doc,
        ["Train Range", "Execution Rule", "Closure Rule"],
        [["L-Z", "Strict sequential execution", "Atlas V5 completes only at 750/750 PASS with evidence committed."]],
    )


def add_train(doc: Document, train: dict) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    tid = train["id"]
    doc.add_heading(f"Train {tid}: {train['name']}", level=1)
    add_table(
        doc,
        ["Field", "Definition"],
        [
            ["Master Objective", train["objective"]],
            ["Primary Routes / Surfaces", train["routes"]],
            ["Dependencies", train["dependencies"]],
            ["Acceptance Target", f"50/50 PASS across WP-{tid}-01 through WP-{tid}-05"],
        ],
    )
    for index, (title, summary) in enumerate(train["wps"], start=1):
        wp_id = f"WP-{tid}-{index:02d}"
        doc.add_heading(f"{wp_id}: {title}", level=2)
        doc.add_paragraph(summary)
        doc.add_heading("Acceptance Criteria", level=3)
        for criterion_index, template in enumerate(CRITERION_TEMPLATES, start=1):
            criterion_id = f"AC-{tid}-{index:02d}-{criterion_index:02d}"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            r = p.add_run(f"{criterion_id}: ")
            r.bold = True
            p.add_run(template)
        doc.add_heading("Required Evidence", level=3)
        for ev in [
            f"EV-CODE_{wp_id}_<YYYYMMDD>_<brief>.*",
            f"EV-TEST_{wp_id}_<YYYYMMDD>_<brief>.*",
            f"EV-ACPT_{wp_id}_<YYYYMMDD>_<brief>.*",
            f"EV-DEPL_{wp_id}_<YYYYMMDD>_<brief>.* when deployment or route behavior is touched",
        ]:
            doc.add_paragraph(ev, style="List Bullet")
    doc.add_heading(f"Train {tid} Closure Definition", level=2)
    doc.add_paragraph(
        f"Train {tid} closes only when WP-{tid}-01 through WP-{tid}-05 each achieve 10/10 PASS, evidence is committed, "
        f"Kimi verifies completeness, and ONX_TRAIN_{tid}_FINAL_VERIFICATION_v1.0.docx is issued."
    )


def add_master_sections(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Part 5: Master Acceptance Matrix", level=1)
    rows = []
    for train in TRAINS:
        tid = train["id"]
        rows.append([f"Train {tid}", f"WP-{tid}-01 to WP-{tid}-05", "50", "Pending execution", "Kimi + Founder"])
    add_table(doc, ["Train", "WP Range", "Criteria", "Initial Status", "Reviewer"], rows)

    doc.add_heading("Part 6: Master Evidence Standard", level=1)
    add_table(
        doc,
        ["Category", "Purpose", "Required Pattern"],
        [
            ["Code", "Implementation proof", "EV-CODE_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Test", "Automated or scripted verification", "EV-TEST_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Security", "Auth, RBAC, or safety proof", "EV-SEC_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Deployment", "Build, staging, release, or rollback proof", "EV-DEPL_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Runtime", "Health, latency, smoke, or monitoring proof", "EV-RUN_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Acceptance", "Checklist and reviewer decision", "EV-ACPT_<WP-ID>_<YYYYMMDD>_<brief>.*"],
            ["Closure", "Train or program closure order", "EV-CLSR_<EP-ID>_<YYYYMMDD>_<brief>.*"],
        ],
    )

    doc.add_heading("Part 7: Master Dependency Graph", level=1)
    doc.add_paragraph("K -> L -> M -> N -> O -> P -> Q -> R -> S -> T -> U -> V -> W -> X -> Y -> Z")
    doc.add_paragraph("No train begins before predecessor closure. No work package begins before predecessor WP closure inside the same train.")

    doc.add_heading("Part 8: Execution Sequencing Rules", level=1)
    for rule in [
        "Execute trains in strict alphabetical order from L through Z.",
        "Execute WPs sequentially from WP-01 through WP-05 inside each train.",
        "Before WP submission, run lint, test, build, and route checks relevant to the WP.",
        "Every acceptance decision must cite concrete evidence and a commit hash.",
        "Any failed criterion blocks WP closure until remediation evidence is committed.",
    ]:
        doc.add_paragraph(rule, style="List Number")

    doc.add_heading("Part 9: Kimi Verification Instructions", level=1)
    for instruction in [
        "Verify every WP has exactly 10 acceptance criteria.",
        "Verify the evidence files exist, are current, and map to the claimed WP.",
        "Score each criterion as PASS or FAIL only.",
        "Reject verbal confirmation without a reviewable artifact.",
        "Escalate any ambiguity, missing evidence, or scope drift to Founder before closure.",
    ]:
        doc.add_paragraph(instruction, style="List Number")

    doc.add_heading("Part 10: Founder Finalization Clause", level=1)
    doc.add_paragraph(
        "Atlas V5 is not complete until all 15 trains, all 75 work packages, and all 750 criteria are scored PASS with evidence "
        "committed to the repository. The authorized completion target is 750/750 PASS."
    )


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_title(doc)
    add_overview(doc)
    for train in TRAINS:
        add_train(doc, train)
    add_master_sections(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
