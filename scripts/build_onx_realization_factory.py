from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs")
FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)
FILL = "E8EEF5"


def set_font(run, size=11, color=INK, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def shade(cell, fill=FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        s = styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def para(doc, text="", style=None, size=11, color=INK, bold=False, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_font(r, size=10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=9.5)
    doc.add_paragraph()
    return t


def add_title(doc):
    para(doc, "ONX SYSTEMS", bold=True, color=BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "ONX Realization Factory", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Permanent execution infrastructure for ATRP-2026-001", color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    table(
        doc,
        ["Document ID", "Version", "Status", "Authority"],
        [["ONX-REALIZATION-FACTORY-2026-001", "1.0", "OPERATIONAL", "Founder -> Atlas -> Kimi -> Husam"]],
        [3100, 1100, 1800, 3360],
    )


def build_factory():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    para(doc, "Executive Operating Declaration", "Heading 1")
    para(doc, "The ONX Realization Factory converts program preparation into a permanent execution system. It keeps planning ahead of execution, keeps issuance as derivation rather than creation, and preserves constitutional control through Atlas V5.0 and CCR governance.")
    table(
        doc,
        ["Factory State", "Value"],
        [["Atlas V5.0", "Frozen"], ["EP-01", "Issued - Active"], ["EP-02", "Pre-Issuance - Ready"], ["EP-03/04/05", "Pre-Issuance - Ready"], ["Master Pipeline", "Active"], ["Husam Continuity", "Never blocked by planning"]],
        [2800, 6560],
    )

    para(doc, "Deliverable 01: EP-03 Pre-Issuance Package", "Heading 1")
    bullets(doc, ["Canonical file: ONX_EP03_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "Compatibility file: ONX_EP03_PRE_ISSUANCE_v1.0.docx", "Scope: User Experience & Interface", "Seven work packages, 56 acceptance criteria, 56 evidence items, transition to EP-04."])

    para(doc, "Deliverable 02: EP-04 Pre-Issuance Package", "Heading 1")
    bullets(doc, ["Canonical file: ONX_EP04_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "Compatibility file: ONX_EP04_PRE_ISSUANCE_v1.0.docx", "Scope: Operations & Deployment", "Six work packages, 48 acceptance criteria, 48 evidence items, transition to EP-05."])

    para(doc, "Deliverable 03: EP-05 Pre-Issuance Package", "Heading 1")
    bullets(doc, ["Canonical file: ONX_EP05_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "Compatibility file: ONX_EP05_PRE_ISSUANCE_v1.0.docx", "Scope: Intelligence & Automation", "Five work packages, 40 acceptance criteria, 40 evidence items, transition to ATRP close."])

    para(doc, "Deliverable 04: Master Execution Pipeline", "Heading 1")
    bullets(doc, ["Canonical file: ONX_MASTER_EXECUTION_PIPELINE_v1.0.docx", "Covers EP-01 through EP-05 sequencing, release gates, evidence gates, execution density, and Husam continuity.", "Recommended pipeline duration: 41-50 weeks.", "Maximum compression range: 33-40 weeks with Founder-authorized compression."])

    para(doc, "Deliverable 05: Master Execution Wave Plan", "Heading 1")
    table(
        doc,
        ["Wave", "Scope", "Activation Criteria", "Exit Criteria"],
        [["Wave 1", "EP-01 foundation", "Founder issuance complete", "EP-02 readiness gate achieved"], ["Wave 2", "EP-02 services", "EP-01 M-08 passed", "EP-03 readiness gate achieved"], ["Wave 3", "EP-03 user experience", "EP-02 contracts frozen", "EP-04 readiness gate achieved"], ["Wave 4", "EP-04 operations", "EP-03 build and compliance artifacts ready", "EP-05 readiness gate achieved"], ["Wave 5", "EP-05 intelligence", "EP-04 production readiness passed", "ATRP close package accepted"]],
        [1100, 2200, 3000, 3060],
    )

    para(doc, "Deliverable 06: Release Train Architecture", "Heading 1")
    table(
        doc,
        ["Train", "Name", "Scope", "Exit"],
        [["Train A", "Foundation", "Architecture, framework, database, API, security, CI/CD", "EP-01 complete"], ["Train B", "Service Runtime", "Runtime host, services, endpoints, jobs, deployment", "EP-02 complete"], ["Train C", "Experience", "Shell, components, dashboards, accessibility, frontend tests", "EP-03 complete"], ["Train D", "Operations", "Infrastructure, monitoring, DR, performance, runbooks", "EP-04 complete"], ["Train E", "Intelligence", "Analytics, automation, AI, recommendations, intelligent monitoring", "EP-05 complete"]],
        [1100, 1900, 4700, 1660],
    )

    para(doc, "Deliverable 07: Maximum Throughput Plan", "Heading 1")
    table(
        doc,
        ["Lever", "Saved Time", "Control"],
        [["Parallel WP branches", "6 weeks", "Only after foundation WP exits"], ["Early evidence capture", "2 weeks", "Registry updated daily"], ["Contract-first handoffs", "3 weeks", "Contract tests required"], ["Continuous docs", "2 weeks", "Runbooks drafted during build"], ["Security pre-review", "2 weeks", "Threat model starts at WP-01"], ["Downstream pre-read", "2 weeks", "Next EP package available before issuance"]],
        [2900, 1500, 4960],
    )

    para(doc, "Deliverable 08: Evidence Factory Standard", "Heading 1")
    bullets(doc, ["Every WP opens with an evidence manifest.", "Every completion claim requires CODE, TEST, DOC, REVIEW, and where applicable SEC/PERF/SCREEN evidence.", "Evidence registry is append-only after WP completion.", "Evidence gates EV-GATE-01 through EV-GATE-05 must pass before release gates close."])
    table(doc, ["Evidence Gate", "Program", "Minimum Verdict"], [["EV-GATE-01", "EP-01", "PASS"], ["EV-GATE-02", "EP-02", "PASS"], ["EV-GATE-03", "EP-03", "PASS"], ["EV-GATE-04", "EP-04", "PASS"], ["EV-GATE-05", "EP-05", "PASS"]], [2200, 2200, 4960])

    para(doc, "Deliverable 09: Issuance Derivation Protocol", "Heading 1")
    numbered(doc, ["Confirm upstream release gate passed.", "Confirm evidence gate passed.", "Confirm Founder approval or delegated authority.", "Derive issuance package from pre-issuance source.", "Replace PRE-ISSUANCE status with ISSUED status.", "Attach start order and reporting cadence.", "Record issuance in document history.", "Notify Husam with canonical package path."])

    para(doc, "Deliverable 10: Husam Never-Blocked Operating Model", "Heading 1")
    bullets(doc, ["Planning is always at least one EP ahead of execution.", "Execution package creation is never on the critical path.", "Next-program dependency maps are available before current-program close.", "CCR pathways are explicit for any deviation.", "Report templates and evidence registry structures are predefined.", "Release train and wave views give alternative navigation when schedule pressure changes.", "No work begins without authority, but no authorized work waits for missing package structure.", "Husam always has the next executable reference ready before activation."])

    para(doc, "Deliverable 11: Factory Governance and Continuous Improvement Loop", "Heading 1")
    table(
        doc,
        ["Cadence", "Owner", "Purpose", "Output"],
        [["Daily", "Husam", "Execution status and blockers", "Daily status report"], ["Weekly", "Kimi", "Evidence and dependency review", "Weekly factory health note"], ["Per WP", "Kimi + Husam", "Completion verification", "WP verdict"], ["Per EP", "Founder + Kimi", "Release gate verdict", "EP close / next EP issue"], ["Post-ATRP", "Founder", "Factory retrospective", "Factory v1.1 improvements"]],
        [1300, 1900, 3300, 2860],
    )

    para(doc, "Appendix A: Document Inventory", "Heading 1")
    table(
        doc,
        ["Document", "Role"],
        [["ONX_EXECUTION_HANDOVER_PACKAGE_EP01_v1.0.docx", "EP-01 issued package"], ["ONX_EP02_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-02 pre-issuance package"], ["ONX_EP03_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-03 pre-issuance package"], ["ONX_EP04_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-04 pre-issuance package"], ["ONX_EP05_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-05 pre-issuance package"], ["ONX_MASTER_EXECUTION_PIPELINE_v1.0.docx", "Unified execution pipeline"], ["ONX_REALIZATION_FACTORY_v1.0.docx", "Permanent factory operating model"]],
        [5000, 4360],
    )
    para(doc, "Appendix B: Activation Checklist", "Heading 1")
    bullets(doc, ["☐ Factory document rendered and verified", "☐ Compatibility filenames created", "☐ EP package structural counts verified", "☐ Master pipeline rendered and verified", "☐ Realization Factory declared operational"])
    para(doc, "Final Declaration", "Heading 1")
    para(doc, "ONX Realization Factory is OPERATIONAL. Execution preparation remains ahead of execution. Husam continuity is active. ATRP-2026-001 proceeds under Atlas V5.0 frozen authority.")
    doc.save(OUT / "ONX_REALIZATION_FACTORY_v1.0.docx")


def make_compatibility_copies():
    pairs = [
        ("ONX_EP03_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "ONX_EP03_PRE_ISSUANCE_v1.0.docx"),
        ("ONX_EP04_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "ONX_EP04_PRE_ISSUANCE_v1.0.docx"),
        ("ONX_EP05_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "ONX_EP05_PRE_ISSUANCE_v1.0.docx"),
    ]
    for src, dst in pairs:
        copyfile(OUT / src, OUT / dst)


if __name__ == "__main__":
    make_compatibility_copies()
    build_factory()
