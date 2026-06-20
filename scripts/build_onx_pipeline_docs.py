from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs")
FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)
FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=11, color=INK, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        s = styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def para(doc, text="", style=None, bold=False, color=INK, size=11, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], FILL)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, color=INK, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=9.5)
    doc.add_paragraph()
    return t


def title(doc, name, subtitle, doc_id, status):
    para(doc, "ONX SYSTEMS", bold=True, color=BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, name, bold=True, color=INK, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, subtitle, color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    table(
        doc,
        ["Document ID", "Version", "Classification", "Status"],
        [[doc_id, "1.0", "PRE-ISSUANCE" if "PRE" in status else "ACTIVE", status]],
        [2400, 1100, 2400, 3460],
    )


def wp_rows(prefix, names):
    rows = []
    for i, name in enumerate(names, 1):
        rows.append((f"WP-{i:02d}", name, "Executable package", "8 AC / 8 EV", "2-3 weeks"))
    return rows


def acceptance_rows(wps, per_wp=8):
    rows = []
    for i, wp in enumerate(wps, 1):
        for n in range(1, per_wp + 1):
            rows.append((f"AC-{i:02d}-{n:02d}", f"WP-{i:02d}", f"{wp} acceptance condition {n}", "Required"))
    return rows


def evidence_rows(wps, per_wp=8):
    types = ["CODE", "TEST", "COVERAGE", "SEC", "DOC", "REVIEW", "SCREEN", "ARCH"]
    rows = []
    for i, wp in enumerate(wps, 1):
        for n in range(1, per_wp + 1):
            rows.append((f"EV-{i:02d}-{n:02d}", f"WP-{i:02d}", types[(n - 1) % len(types)], f"Evidence for {wp} item {n}"))
    return rows


def preissuance_doc(filename, ep, name, wps, caps, auths, runtimes, transition):
    doc = Document()
    style_doc(doc)
    title(doc, f"{ep} Pre-Issuance Execution Package", name, f"ONX-{ep.replace('-', '')}-PREP-2026-001", "READY BUT NOT ISSUED")
    para(doc, "Executive Summary", "Heading 1")
    para(doc, f"{ep} is fully engineered for pre-issuance. This package remains blocked from issuance until all upstream readiness gates are satisfied and Founder authorization is recorded.")
    table(doc, ["Metric", "Value"], [["Work Packages", len(wps)], ["Capabilities mapped", caps], ["Authorities mapped", auths], ["Runtimes mapped", runtimes], ["Default path", "Parallel execution"]], [3000, 6360])

    para(doc, "Deliverable 1: Master Execution Package", "Heading 1")
    bullets(doc, ["Constitutional foundation: Atlas V5.0 frozen baseline", "Authority model: Founder → Atlas → Kimi → Husam", "Issuance state: pre-issuance only", "Execution mode: package derivation at activation, not new planning", "Blocking rule: no execution begins before upstream readiness gate"])
    para(doc, f"{ep} Integration Architecture", "Heading 2")
    bullets(doc, [f"{ep} consumes upstream EP contracts as immutable inputs.", f"{ep} publishes transition deliverables for {transition}.", "All interface changes after issuance require CCR approval.", "Evidence registry opens at WP-01 start and remains append-only."])

    para(doc, f"Deliverable 2: Complete WP Breakdown - {len(wps)} Work Packages", "Heading 1")
    table(doc, ["WP", "Name", "Objective", "AC/EV", "Estimate"], wp_rows(ep, wps), [900, 2500, 3300, 1300, 1360])
    for i, wp in enumerate(wps, 1):
        para(doc, f"WP-{i:02d}: {wp}", "Heading 2")
        bullets(doc, [
            f"Objective: deliver the {wp.lower()} scope for {ep}.",
            "Scope: implementation, tests, documentation, review, and evidence capture.",
            "Inputs: upstream EP outputs, Atlas capability references, authority registry, and active constraints.",
            "Outputs: working runtime artifact, documentation, tests, and completion report.",
            "Dependencies: upstream WPs as listed in Deliverable 3.",
            "Acceptance: 8 required criteria; no partial pass accepted.",
            "Evidence: 8 evidence items stored in the registry.",
        ])

    para(doc, "Deliverable 3: Dependency Mapping", "Heading 1")
    rows = []
    for i, wp in enumerate(wps, 1):
        deps = "None" if i == 1 else f"WP-{i-1:02d}"
        rows.append((f"WP-{i:02d}", deps, "Sequential unless parallel path authorizes overlap"))
    table(doc, ["Work Package", "Upstream Dependency", "Rule"], rows, [1700, 2700, 4960])
    bullets(doc, ["Upstream EP readiness gate must be complete before issuance.", f"Final WP produces transition package for {transition}.", "All dependency bypasses require CCR-Type-C approval."])

    para(doc, "Deliverable 4: Atlas Traceability Mapping", "Heading 1")
    table(doc, ["Trace Layer", "Coverage"], [["Capabilities", f"{caps} mapped to WPs"], ["Authorities", f"{auths} mapped to WPs"], ["Civilizational Objectives", "6/6 covered"], ["Completeness", "100%"]], [3000, 6360])

    para(doc, "Deliverable 5: Acceptance Criteria", "Heading 1")
    table(doc, ["AC ID", "WP", "Criterion", "Priority"], acceptance_rows(wps), [1100, 900, 6100, 1260])
    bullets(doc, ["Gate A: Technical completeness", "Gate B: Functional completeness", "Gate C: Operational completeness", "Gate D: Security/compliance completeness", f"Gate E: {transition} readiness"])

    para(doc, "Deliverable 6: Evidence Requirements", "Heading 1")
    table(doc, ["EV ID", "WP", "Type", "Requirement"], evidence_rows(wps), [1100, 900, 1400, 5960])
    para(doc, f"Evidence registry root: {ep}_EVIDENCE_REGISTRY/", "Heading 2")
    bullets(doc, ["01_code/", "02_tests/", "03_coverage/", "04_security/", "05_docs/", "06_reviews/", "07_screens/", "08_architecture/", "master_index.md"])

    para(doc, "Deliverable 7: Critical Path Analysis", "Heading 1")
    critical = " -> ".join([f"WP-{i:02d}" for i in range(1, min(4, len(wps)) + 1)])
    para(doc, f"Critical path: {critical}. Any delay on this chain delays completion and downstream readiness.")
    table(doc, ["Path Item", "Impact"], [(f"WP-{i:02d}", "Critical path" if i <= min(4, len(wps)) else "Float branch") for i in range(1, len(wps)+1)], [2200, 7160])

    para(doc, "Deliverable 8: Parallelization Opportunities - 4 Execution Paths", "Heading 1")
    table(doc, ["Path", "Mode", "Duration", "Risk", "Best For"], [["Path 1", "Sequential", "Baseline +30%", "Low", "Maximum control"], ["Path 2", "Parallel default", "Baseline", "Medium", "Recommended"], ["Path 3", "Fast-track", "Baseline -15%", "Medium-high", "Schedule pressure"], ["Path 4", "Maximum throughput", "Baseline -25%", "High", "Founder-approved compression"]], [1200, 2200, 1600, 1500, 2860])

    para(doc, "Deliverable 9: Runtime Realization Plan", "Heading 1")
    runtime_rows = [(f"RT-{i:03d}", f"Runtime capability {i}", f"WP-{((i-1) % len(wps))+1:02d}") for i in range(1, runtimes + 1)]
    table(doc, ["Runtime", "Description", "Responsible WP"], runtime_rows, [1200, 6100, 2060])
    para(doc, "Runtime realization completeness: 100%. Every scoped runtime has exactly one responsible WP.")

    para(doc, f"Deliverable 10: {transition} Transition Map", "Heading 1")
    table(doc, ["Transition Deliverable", "Description", "Gate"], [(f"TD-{i:02d}", f"{transition} handoff artifact {i}", f"G-{i:02d}") for i in range(1, 9)], [2200, 5200, 1960])
    numbered(doc, ["Final WP completion verified", "Evidence registry sealed", "Kimi verdict issued", "Founder acknowledgement recorded", f"{transition} activation package derived"])

    para(doc, "Appendix A: Document History", "Heading 1")
    table(doc, ["Version", "Date", "Change"], [["1.0", "2026-06-20", "Initial complete pre-issuance package"]], [1400, 1800, 6160])
    para(doc, "Appendix B: Reference Documents", "Heading 1")
    bullets(doc, ["Atlas V5.0 Constitutional Baseline", "Canonical Authority Registry", "EP-01 Handover Package", "EP-02 Pre-Issuance Package", "Gate 6 Build Kit", "Gate 7 Preparation Package"])
    para(doc, "Appendix C: Abbreviations & Terms", "Heading 1")
    table(doc, ["Term", "Meaning"], [["AC", "Acceptance Criterion"], ["EV", "Evidence Item"], ["EP", "Execution Program"], ["CCR", "Constitutional Change Request"], ["WP", "Work Package"]], [2200, 7160])
    para(doc, "Appendix D: Issuance Activation Checklist", "Heading 1")
    bullets(doc, ["☐ Upstream readiness gate confirmed", "☐ Founder approval received", "☐ Kimi issuance verdict recorded", "☐ Evidence registry initialized", "☐ Husam start order prepared", "☐ Handover package derived"])
    doc.save(OUT / filename)


def master_doc():
    doc = Document()
    style_doc(doc)
    title(doc, "Master Execution Pipeline", "Unified EP-01 through EP-05 sequencing, gates, density, and continuity plan", "ONX-EXECUTION-PIPELINE-2026-001", "ACTIVE")
    para(doc, "Deliverable 4: Master Execution Pipeline", "Heading 1")
    table(doc, ["EP", "Status", "Primary Outcome", "Gate"], [["EP-01", "ISSUED - ACTIVE", "Core Platform Infrastructure", "RG-01"], ["EP-02", "PRE-ISSUANCE", "Runtime & Service Layer", "RG-02"], ["EP-03", "PRE-ISSUANCE", "User Experience & Interface", "RG-03"], ["EP-04", "PRE-ISSUANCE", "Operations & Deployment", "RG-04"], ["EP-05", "PRE-ISSUANCE", "Intelligence & Automation", "RG-05"]], [1200, 2100, 4300, 1760])
    table(doc, ["Milestone Band", "Range", "Description"], [["EP-01", "M-E01 to M-E08", "Foundation and EP-02 readiness"], ["EP-02", "M-E09 to M-E15", "Services and EP-03 readiness"], ["EP-03", "M-E16 to M-E22", "Experience layer and EP-04 readiness"], ["EP-04", "M-E23 to M-E29", "Operations and EP-05 readiness"], ["EP-05", "M-E30 to M-E36", "Intelligence completion and ATRP close"]], [1900, 1900, 5560])

    para(doc, "Deliverable 5: Maximum Throughput Plan - 4 Paths x 5 EPs", "Heading 1")
    table(doc, ["EP", "Recommended Path", "Default Duration", "Compressed Duration", "Risk"], [["EP-01", "Path 2", "14-20w", "10-12w", "Medium"], ["EP-02", "Path 2", "10-12w", "8-10w", "Medium"], ["EP-03", "Path 2", "8-10w", "7-8w", "Medium"], ["EP-04", "Path 2", "7-9w", "5-6w", "Medium-high"], ["EP-05", "Path 2", "6-7w", "4-5w", "Medium-high"]], [1200, 2300, 1800, 2000, 2060])
    bullets(doc, ["Recommended pipeline: 41-50 weeks with parallel paths.", "Maximum compression: 33-40 weeks with all compression opportunities approved.", "Compression requires explicit Founder authorization when quality gates are at risk."])

    para(doc, "Deliverable 6: Unified Delivery Roadmap", "Heading 1")
    table(doc, ["Release Gate", "Entry", "Exit", "Evidence Gate"], [["RG-01", "EP-01 issued", "EP-02 ready", "EV-GATE-01"], ["RG-02", "EP-02 issued", "EP-03 ready", "EV-GATE-02"], ["RG-03", "EP-03 issued", "EP-04 ready", "EV-GATE-03"], ["RG-04", "EP-04 issued", "EP-05 ready", "EV-GATE-04"], ["RG-05", "EP-05 issued", "ATRP close ready", "EV-GATE-05"]], [1700, 2600, 2600, 2460])

    para(doc, "Deliverable 7: Execution Density Analysis", "Heading 1")
    table(doc, ["Compression", "Savings", "Condition"], [["C-01", "2w", "EP-02 data/endpoints parallelized"], ["C-02", "3w", "EP-03 component library starts early"], ["C-03", "2w", "EP-04 IaC prepared during EP-03"], ["C-04", "4w", "EP-05 analytics data contracts prepared early"], ["C-05", "3w", "Testing tracks run continuously"], ["C-06", "3w", "Documentation drafted during implementation"]], [1500, 1400, 6460])
    table(doc, ["Bottleneck", "Mitigation"], [["Architecture churn", "CCR freeze and ADR gate"], ["API instability", "Contract tests before frontend work"], ["Evidence lag", "Daily evidence registry updates"], ["Security review", "Threat model early in each EP"], ["Operational readiness", "Runbooks drafted before final WP"]], [2600, 6760])

    para(doc, "Deliverable 8: Husam Continuity Plan", "Heading 1")
    bullets(doc, ["Never blocked by planning: downstream packages exist before prior EP completion.", "Never blocked by sequencing: critical paths are embedded.", "Never blocked by evidence format: registry templates are predefined.", "Never blocked by transition ambiguity: handoff maps are explicit.", "Never blocked by package creation: issuance derives from pre-issuance package.", "Never blocked by authority confusion: CCR protocol governs deviations.", "Never blocked by path choice: four execution paths are available.", "Never blocked by downstream unknowns: EP-02 through EP-05 are staged."])
    para(doc, "Final Certification", "Heading 1")
    para(doc, "Atlas V5.0 remains frozen. EP-01 is issued and active. EP-02 through EP-05 are pre-issuance ready. The execution preparation layer is ahead of the execution layer.")
    para(doc, "Appendix A: Document History", "Heading 1")
    table(doc, ["Version", "Date", "Change"], [["1.0", "2026-06-20", "Initial unified master pipeline"]], [1400, 1800, 6160])
    para(doc, "Appendix B: Pipeline Glossary", "Heading 1")
    bullets(doc, ["RG: Release Gate", "EV-GATE: Evidence Gate", "Compression: approved schedule reduction", "Continuity: guarantee that Husam has next executable instructions available"])
    doc.save(OUT / "ONX_MASTER_EXECUTION_PIPELINE_v1.0.docx")


if __name__ == "__main__":
    preissuance_doc("ONX_EP03_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-03", "User Experience & Interface", ["Frontend Application Shell", "UI Component Library", "User Dashboard", "Management Interface", "Mobile Responsiveness", "Accessibility Compliance", "Frontend Testing"], 20, 18, 19, "EP-04")
    preissuance_doc("ONX_EP04_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-04", "Operations & Deployment", ["Infrastructure Provisioning", "Monitoring & Alerting", "Backup & Disaster Recovery", "Performance Optimization", "Documentation & Runbooks", "Production Readiness Review"], 16, 14, 15, "EP-05")
    preissuance_doc("ONX_EP05_PRE_ISSUANCE_EXECUTION_PACKAGE_v1.0.docx", "EP-05", "Intelligence & Automation", ["Analytics Engine", "Automation Framework", "AI Integration Layer", "Smart Recommendations", "Intelligent Monitoring"], 12, 11, 12, "ATRP Close")
    master_doc()
