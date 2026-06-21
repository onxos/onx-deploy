from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/ONX_ATRP_2026_003_EP01_TECHNICAL_BASELINE_VALIDATION_v1.0.docx")
FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)
FILL = "E8EEF5"
RISK = "FCE4D6"
PASS = "E2F0D9"
WARN = "FFF2CC"


def set_font(run, size=11, color=INK, bold=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        s = styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def para(doc, text="", style=None, size=11, color=INK, bold=False, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def table(doc, headers, rows, widths=None, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], FILL)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_font(r, size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                width(cells[i], widths[i])
            if status_col is not None and i == status_col:
                text = str(value).upper()
                shade(cells[i], PASS if "PASS" in text or "READY" == text else WARN if "PARTIAL" in text else RISK)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 16 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=9)
    doc.add_paragraph()
    return t


def build():
    doc = Document()
    style_doc(doc)
    para(doc, "ONX SYSTEMS", bold=True, color=BLUE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "ONX-ATRP-2026-003: EP-01 Technical Baseline Validation", bold=True, size=21, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Corrected validation package for the actual Next.js + tRPC baseline", color=MUTED, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    table(doc, ["Document ID", "Version", "Date", "Gate Decision"], [["ONX-ATRP-2026-003", "1.0", "2026-06-20", "REMEDIATION REQUIRED"]], [2300, 1100, 1700, 4260], status_col=3)

    para(doc, "Executive Verdict", "Heading 1")
    para(doc, "EP-01 is verified as a functional Next.js + tRPC application. The baseline builds, lints, starts, serves pages, and executes public tRPC query/mutation flows. However, staging deployment is not fully authorized until operational remediation is complete.")
    bullets(doc, [
        "Corrected stack: Next.js 16 App Router, React 19, tRPC v11, Drizzle ORM, Neon/Postgres.",
        "Rejected stale claims: Hono backend, Vite frontend, SQLite WAL, Docker Compose, nginx, and separate port 3001 are not present in this repo.",
        "Current gate: REMEDIATION REQUIRED before staging deployment readiness.",
    ])

    para(doc, "D01: Technical Baseline Report", "Heading 1")
    table(doc, ["Item", "Observed Baseline"], [
        ["Commit", "dd8802de5c25b178b8090f3557aa57401523c92f"],
        ["Branch", "main"],
        ["Runtime", "Next.js 16.2.9 production build"],
        ["Source files", "94 TypeScript/TSX files under src"],
        ["Source LOC", "4,321 lines under src"],
        ["tRPC routers", "6 routers plus root router"],
        ["tRPC procedures", "28 procedure definitions observed"],
        ["Database schema", "9 Postgres tables across Drizzle schema modules"],
    ], [2600, 6760])

    para(doc, "D02: Environment Report", "Heading 1")
    table(doc, ["Environment Item", "Observed Value"], [
        ["Node", "v24.17.0"],
        ["npm", "11.13.0"],
        ["Bun", "1.3.14"],
        ["OS", "Darwin arm64"],
        ["Database dependency", "DATABASE_URL required"],
        ["Admin dependency", "ADMIN_ACCESS_TOKEN required"],
        ["Auth dependency", "BETTER_AUTH_SECRET and BETTER_AUTH_URL configured in env"],
    ], [2600, 6760])
    bullets(doc, [
        "Local environment is not Debian 12 and does not match the stale validation text.",
        "Vercel CLI remains outdated in the local machine context; upgrade recommended before deployment work.",
        "No Dockerfiles, docker-compose, nginx config, or CI workflow files were found.",
    ])

    para(doc, "D03: Deployment Readiness Assessment", "Heading 1")
    table(doc, ["Readiness Area", "Status", "Evidence"], [
        ["Production build", "READY", "bun run build passed"],
        ["Static/dynamic route compilation", "READY", "16 app routes generated; /api/trpc dynamic"],
        ["Lint/format gate", "READY", "bun run lint passed; Biome checked 101 files"],
        ["Runtime server", "READY", "next start served localhost:3000 in prior smoke test"],
        ["tRPC public query", "READY", "civilization.listArticles returned seeded article data"],
        ["tRPC public mutation", "READY", "analytics.track inserted row and returned HTTP 200"],
        ["Admin auth gate", "PARTIALLY_READY", "Token-based gate exists; Better Auth sessions not wired into protectedProcedure"],
        ["Database migration/deploy workflow", "PARTIALLY_READY", "Drizzle config exists; automated deployment migration gate missing"],
        ["CI/CD pipeline", "NOT_READY", "No .github workflow or equivalent CI config found"],
        ["Backup/restore automation", "NOT_READY", "No database backup or restore automation found"],
    ], [2600, 1800, 4960], status_col=1)

    para(doc, "D04: Known Issues Register", "Heading 1")
    table(doc, ["ID", "Severity", "Issue", "Remediation"], [
        ["ISSUE-001", "HIGH", "Formal CI/CD pipeline missing", "Add GitHub Actions or Vercel checks for build, lint, and migration dry-run"],
        ["ISSUE-002", "HIGH", "Backup/restore automation missing", "Define Neon backup policy and restore drill evidence"],
        ["ISSUE-003", "HIGH", "Protected procedures use temporary admin token gate", "Wire Better Auth sessions/roles into tRPC context"],
        ["ISSUE-004", "MEDIUM", "Deployment config absent", "Add vercel.ts or documented Vercel project configuration"],
        ["ISSUE-005", "MEDIUM", "No automated test suite present", "Add focused tests for routers and critical UI flows"],
        ["ISSUE-006", "MEDIUM", "No explicit health endpoint", "Add /api/health or equivalent operational probe"],
        ["ISSUE-007", "MEDIUM", "Database migrations not validated in CI", "Run drizzle-kit checks in deployment pipeline"],
        ["ISSUE-008", "LOW", "Outdated local Vercel CLI", "Upgrade to latest Vercel CLI before deployment"],
        ["ISSUE-009", "LOW", "Documented stale stack claims existed", "Correct validation docs to Next.js + tRPC baseline"],
    ], [1100, 1300, 3900, 3060])

    para(doc, "D05: Technical Evidence Package", "Heading 1")
    table(doc, ["Check", "Result", "Evidence"], [
        ["Build", "PASS", "next build compiled successfully"],
        ["TypeScript", "PASS", "TypeScript phase completed during next build"],
        ["Lint", "PASS", "Biome checked 101 files with no fixes applied"],
        ["Route smoke", "PASS", "Seven public routes returned HTTP 200 in prior runtime test"],
        ["tRPC query", "PASS", "civilization.listArticles returned article data"],
        ["tRPC mutation", "PASS", "analytics.track wrote row id 201 during smoke test"],
        ["Browser smoke", "PASS", "Homepage and Knowledge page rendered; no captured console errors"],
        ["Integrity hashes", "PASS", "SHA-256 hashes captured for source files"],
    ], [1800, 1200, 6360], status_col=1)
    para(doc, "Evidence sample hashes", "Heading 2")
    table(doc, ["File", "SHA-256"], [
        ["src/app/page.tsx", "94a270ad4eb1916cd29531dbe9f444187ab00c021ffe0e329f7529a9fafcbc04"],
        ["src/app/api/trpc/[trpc]/route.ts", "0f7c1b5fdf81526932ee66443671c764c55f372079c8e397be978d28bbab5f40"],
        ["src/server/api/root.ts", "113d0061a5f7d7c4ab5e71e88ae6305ea419725394fa4536d7dc92eac9bea303"],
        ["src/server/api/trpc.ts", "c14f63713f7998d0f23b5b2f1078a80380751ae43c90c21ceec558b0fc6e62b0"],
        ["src/server/db/index.ts", "950e6b1d5f3b4e3a5e857b05208e793177f64b07550f566f1215d35c97e9a10d"],
    ], [3600, 5760])

    para(doc, "D06: Next Execution Gate", "Heading 1")
    table(doc, ["Gate", "Decision", "Required Before Staging"], [
        ["EP-01 Technical Baseline", "FUNCTIONAL", "No blocker to local runtime validation"],
        ["Staging Deployment", "REMEDIATION REQUIRED", "CI/CD, backup/restore, auth hardening, deployment config"],
        ["Production Deployment", "NOT AUTHORIZED", "Requires staging pass plus Founder approval"],
    ], [2600, 2500, 4260], status_col=1)
    bullets(doc, [
        "Start remediation with CI/CD and deployment config; these unblock repeatable staging checks.",
        "Next, harden admin/session auth and add health/migration gates.",
        "Finally, document backup/restore evidence and add focused automated tests.",
    ])

    para(doc, "Final Classification", "Heading 1")
    para(doc, "EP-01 is technically verified as a working Next.js + tRPC baseline. The correct gate classification is REMEDIATION REQUIRED for staging deployment readiness, not because the app fails to run, but because operational deployment controls are incomplete.")
    doc.save(OUT)


if __name__ == "__main__":
    build()
