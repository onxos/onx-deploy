"""
Generates ONX_ATLAS_V5_COMPLETE_EXECUTION_PACKAGE_v1.1.docx
Changes from v1.0:
  - Authority: removed Founder Directive, replaced with Kimi Review / Husam Yildirim
  - Date: 2026-06-24
  - Version: v1.1
  - Acceptance matrix: Trains L and M → CLOSED / PASSED
  - Train N → ACTIVE (current)
  - Trains O-Z → Pending execution
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/ONX_ATLAS_V5_COMPLETE_EXECUTION_PACKAGE_v1.1.docx"

TRAIN_STATUS = {
    "L": ("CLOSED", "Analytics, Reporting & Visualization Layer"),
    "M": ("CLOSED", "Outcome, Flourishing & Recognition Workspace"),
    "N": ("ACTIVE", "Evolution, Review & Continuous Improvement Layer"),
    "O": ("PENDING", "Knowledge Corpus Expansion & Editorial Governance"),
    "P": ("PENDING", "Titan Persona Governance & Conversational Operations"),
    "Q": ("PENDING", "Institution Administration & Role Operations"),
    "R": ("PENDING", "Security Hardening & Compliance Evidence"),
    "S": ("PENDING", "Runtime Reliability, Health & Observability"),
    "T": ("PENDING", "Staging, Deployment & Release Governance"),
    "U": ("PENDING", "Data Integrity, Migration & Backup Assurance"),
    "V": ("PENDING", "Performance, Load & Scalability Readiness"),
    "W": ("PENDING", "User Acceptance, Training & Operational Enablement"),
    "X": ("PENDING", "Founder Review, Kimi Verification & Final Acceptance"),
    "Y": ("PENDING", "Production Launch, Monitoring & Closure"),
    "Z": ("PENDING", "Atlas V5 Completion Assurance & Handover"),
}

TRAINS = {
    "L": {
        "objective": "Build comprehensive dashboards, charts, metrics, and reporting surfaces across the Dream-to-Evolution lifecycle.",
        "routes": "/dashboard, /analytics, /reports",
        "deps": "Train K lifecycle data, Train I data table and stat-card primitives, Train H layout shell",
        "wps": [
            ("WP-L-01", "Dashboard Overview", "Implement executive KPI cards for dreams, potentials, goals, active execution, and completion velocity."),
            ("WP-L-02", "Dream Conversion Funnel", "Visualize Dream to Potential to Goal to Milestone to Completion conversion with stage-level rates."),
            ("WP-L-03", "Goal Progress Charts", "Add line, bar, and circular progress charts for goal completion, milestone burnup, and trend comparison."),
            ("WP-L-04", "Execution Velocity & Trends", "Track tasks completed per interval with trend indicators, variance notes, and operator filters."),
            ("WP-L-05", "Report Generation & Export", "Generate reviewable PDF and CSV reports with selectable periods, metrics, and evidence references."),
        ],
    },
    "M": {
        "objective": "Implement outcome capture, lessons learned, achievement recognition, and flourishing signals after goals complete.",
        "routes": "/outcomes, /flourishing, /recognition",
        "deps": "Train K goals and execution records, Train L analytics baseline",
        "wps": [
            ("WP-M-01", "Outcome Recording", "Capture completed goal outcomes with result summaries, evidence links, and stakeholder impact."),
            ("WP-M-02", "Lessons Learned", "Record lessons, decisions, and improvement notes linked to completed execution work."),
            ("WP-M-03", "Achievement Badges", "Issue recognition badges for milestones, streaks, completion quality, and contribution patterns."),
            ("WP-M-04", "Flourishing Timeline", "Render a personal and institutional growth timeline across outcomes and recognitions."),
            ("WP-M-05", "Outcome Review Dashboard", "Summarize outcomes, lessons, recognitions, and completion quality for leadership review."),
        ],
    },
    "N": {
        "objective": "Convert completed outcomes into structured improvement loops, retrospectives, and next-cycle recommendations.",
        "routes": "/evolution, /retrospectives",
        "deps": "Train M outcomes and lessons, Train L reports",
        "wps": [
            ("WP-N-01", "Retrospective Workspace", "Create guided retrospectives for completed goals and trains with owner, finding, and action fields."),
            ("WP-N-02", "Improvement Backlog", "Maintain a backlog of improvement items with priority, evidence, assignment, and status."),
            ("WP-N-03", "Pattern Detection", "Surface recurring lessons, blockers, and strengths from prior outcomes using rule-based summaries."),
            ("WP-N-04", "Recommendation Review", "Convert lessons into proposed next actions with approval, rejection, and deferral states."),
            ("WP-N-05", "Evolution Closure Report", "Generate closure reports that bind outcomes, lessons, improvements, and next-cycle actions."),
        ],
    },
    "O": {
        "objective": "Expand the knowledge corpus while adding editorial controls for quality, provenance, publication, and review.",
        "routes": "/knowledge, /admin/knowledge",
        "deps": "Existing corpus ingestion, admin roles, evidence standard",
        "wps": [
            ("WP-O-01", "Corpus Expansion Workflow", "Add submission, draft, review, approved, and published states for knowledge articles."),
            ("WP-O-02", "Editorial Review Queue", "Build an admin queue for review assignment, comments, rejection reasons, and approval history."),
            ("WP-O-03", "Provenance Metadata", "Capture source, owner, date, confidence, category, and evidence references for each article."),
            ("WP-O-04", "Publication Controls", "Support scheduled publishing, unpublishing, archival, and public visibility checks."),
            ("WP-O-05", "Corpus Quality Audit", "Report missing metadata, stale articles, duplicate topics, and evidence gaps."),
        ],
    },
    "P": {
        "objective": "Move Titan persona operations from static definitions toward governed records, auditability, and safe response operations.",
        "routes": "/admin/titans, /titan-conclave, /ask",
        "deps": "Titan database schema, existing hard-coded personas, Gate 6 conversation layer",
        "wps": [
            ("WP-P-01", "Persona Registry Admin", "Build governed admin editing for Titan domain, tone, prompt, traits, and activation state."),
            ("WP-P-02", "Persona Versioning", "Persist persona versions with change author, reason, approval status, and rollback metadata."),
            ("WP-P-03", "Conversation Audit Trail", "Expose searchable conversation records with persona, source, prompt, response, and timestamp."),
            ("WP-P-04", "Safety Guardrails", "Add refusal, escalation, and confidence rules for unsafe or out-of-domain Titan responses."),
            ("WP-P-05", "Titan Operations Report", "Summarize usage, retrieval quality, escalations, and persona changes for Kimi review."),
        ],
    },
    "Q": {
        "objective": "Harden institution management, role assignment, membership workflows, and admin evidence controls.",
        "routes": "/admin, /institution, /settings",
        "deps": "Better Auth, RBAC helpers, admin UI components",
        "wps": [
            ("WP-Q-01", "Member Lifecycle", "Implement invite, activate, suspend, reactivate, and remove flows for institution members."),
            ("WP-Q-02", "Role Assignment Controls", "Add founder/admin/operator role management with permission checks and audit events."),
            ("WP-Q-03", "Access Review", "Create periodic role review reports with stale access, high-privilege accounts, and exceptions."),
            ("WP-Q-04", "Admin Activity Log", "Record administrative actions with actor, subject, timestamp, route, and result."),
            ("WP-Q-05", "Institution Settings", "Provide governed institution profile, policy, notification, and operational setting controls."),
        ],
    },
    "R": {
        "objective": "Strengthen authentication, authorization, session controls, audit evidence, and compliance readiness.",
        "routes": "/login, /settings/security, /admin/security",
        "deps": "Better Auth, RBAC, route protection, evidence standard",
        "wps": [
            ("WP-R-01", "Session Security", "Add session inventory, revocation, expiration visibility, and secure logout coverage."),
            ("WP-R-02", "Password Policy", "Enforce password strength, reset safeguards, rate limits, and failure evidence."),
            ("WP-R-03", "Authorization Tests", "Expand pass and fail authorization tests across protected routes and procedures."),
            ("WP-R-04", "Security Evidence Pack", "Produce security evidence files for auth flows, role checks, and access denial."),
            ("WP-R-05", "Compliance Review Dashboard", "Expose compliance status, open findings, critical gaps, and closure evidence."),
        ],
    },
    "S": {
        "objective": "Add production-grade health, readiness, metrics, logging, and incident visibility.",
        "routes": "/api/health, /api/health/ready, /api/metrics, /admin/operations",
        "deps": "Existing health endpoints, deployment scripts, monitoring docs",
        "wps": [
            ("WP-S-01", "Health Endpoint Expansion", "Validate app, database, auth, tRPC, and corpus readiness with clear status payloads."),
            ("WP-S-02", "Metrics Surface", "Emit operational metrics for request health, latency, failures, and domain workflows."),
            ("WP-S-03", "Structured Logging", "Standardize logs with event type, severity, request context, actor, and correlation id."),
            ("WP-S-04", "Incident Workflow", "Connect alerts, escalation, incident notes, owner assignment, and closure evidence."),
            ("WP-S-05", "Reliability Report", "Generate uptime, latency, error, smoke test, and rollback readiness reports."),
        ],
    },
    "T": {
        "objective": "Formalize staging validation, deployment evidence, rollback readiness, and release approvals.",
        "routes": "/admin/releases, deployment scripts",
        "deps": "Build pipeline, staging smoke tests, evidence folders",
        "wps": [
            ("WP-T-01", "Release Checklist", "Create a governed checklist for build, lint, smoke, evidence, approval, and rollout readiness."),
            ("WP-T-02", "Staging Smoke Automation", "Expand smoke tests for public, protected, admin, API, health, and auth flows."),
            ("WP-T-03", "Rollback Validation", "Document and test rollback steps with owner, timing, command, and evidence output."),
            ("WP-T-04", "Deployment Evidence Builder", "Collect deployment logs, URLs, health checks, commit hashes, and approval notes."),
            ("WP-T-05", "Release Certification", "Generate final readiness certification for Kimi approval."),
        ],
    },
    "U": {
        "objective": "Protect application data through integrity checks, migration discipline, backup validation, and recovery evidence.",
        "routes": "/admin/data, database scripts",
        "deps": "Drizzle schema, Neon database, migration scripts",
        "wps": [
            ("WP-U-01", "Schema Integrity Check", "Verify table prefixes, required indexes, relations, and schema drift evidence."),
            ("WP-U-02", "Migration Discipline", "Add migration generation, review, execution, and rollback documentation gates."),
            ("WP-U-03", "Backup Verification", "Define backup presence, restore rehearsal, retention, and recovery-point evidence."),
            ("WP-U-04", "Data Quality Reports", "Find orphaned records, invalid states, missing owners, and inconsistent lifecycle links."),
            ("WP-U-05", "Recovery Runbook", "Create database recovery procedures with roles, timing, commands, and validation outputs."),
        ],
    },
    "V": {
        "objective": "Demonstrate acceptable performance under realistic load and document scale constraints.",
        "routes": "All critical user and API routes",
        "deps": "Runtime metrics, staging deployment, analytics routes",
        "wps": [
            ("WP-V-01", "Performance Budget", "Define budgets for page load, API latency, DB query duration, and bundle growth."),
            ("WP-V-02", "Load Test Harness", "Run repeatable concurrent-user tests for public, authenticated, admin, and API journeys."),
            ("WP-V-03", "Database Query Review", "Identify slow queries, missing indexes, N+1 risks, and pagination defects."),
            ("WP-V-04", "Frontend Performance Audit", "Review bundle cost, render stability, layout shifts, and responsive performance."),
            ("WP-V-05", "Scalability Report", "Compile load results, bottlenecks, remediations, residual risks, and sign-off evidence."),
        ],
    },
    "W": {
        "objective": "Prepare users and operators with acceptance scripts, training materials, support workflows, and adoption evidence.",
        "routes": "/help, /support, training docs",
        "deps": "Completed feature trains, evidence standard, support ownership docs",
        "wps": [
            ("WP-W-01", "UAT Script Library", "Create acceptance scripts for core user, admin, Titan, knowledge, and reporting workflows."),
            ("WP-W-02", "Operator Training Guide", "Produce training guidance for daily operation, evidence collection, and issue handling."),
            ("WP-W-03", "Support Workflow", "Define intake, triage, escalation, resolution, and closure for support requests."),
            ("WP-W-04", "Adoption Dashboard", "Track training completion, UAT status, open blockers, and user feedback."),
            ("WP-W-05", "Enablement Certification", "Issue readiness certification for users, operators, and administrators."),
        ],
    },
    "X": {
        "objective": "Package the final review flow so Kimi can verify completion without ambiguity.",
        "routes": "/admin/acceptance, evidence folders",
        "deps": "All prior train evidence, D02/D04/D05/D06 authority workflow",
        "wps": [
            ("WP-X-01", "Acceptance Matrix Finalization", "Create a complete matrix for train, WP, criterion, evidence, reviewer, and result."),
            ("WP-X-02", "Kimi Verification Workspace", "Provide review queue, scoring, comments, exceptions, and sign-off capture."),
            ("WP-X-03", "Review Packet", "Compile executive summary, evidence index, unresolved risks, and decision prompts."),
            ("WP-X-04", "Exception Management", "Track rejected items, remediation owners, due dates, retest evidence, and closure status."),
            ("WP-X-05", "Final Acceptance Order", "Generate D06-style final order with PASS/FAIL status and authorized next steps."),
        ],
    },
    "Y": {
        "objective": "Execute production launch readiness, monitor launch behavior, close Atlas V5, and preserve evidence.",
        "routes": "Production launch surfaces and monitoring endpoints",
        "deps": "Acceptance verification, release certification, runtime monitoring",
        "wps": [
            ("WP-Y-01", "Production Launch Checklist", "Confirm domains, env vars, migrations, health, rollback, access, and approvals."),
            ("WP-Y-02", "Launch Communications", "Prepare launch notice, stakeholder update, support window, and escalation contacts."),
            ("WP-Y-03", "Post-Launch Monitoring", "Monitor health, latency, errors, auth, data writes, and user journey success."),
            ("WP-Y-04", "Closure Evidence Archive", "Commit complete evidence index, retention notes, release hashes, and approvals."),
            ("WP-Y-05", "Atlas V5 Closure", "Issue Atlas V5 completion status, final metrics, owner sign-off, and post-launch watch plan."),
        ],
    },
    "Z": {
        "objective": "Close the remaining 15-train sequence with final assurance, operating handover, and long-term stewardship controls.",
        "routes": "/admin/closure, /admin/handover, repository docs",
        "deps": "Trains L through Y, production launch monitoring, evidence archive",
        "wps": [
            ("WP-Z-01", "Completion Integrity Audit", "Verify all 75 WPs, 750 criteria, evidence records, and closure documents are present."),
            ("WP-Z-02", "Operational Handover", "Create role-specific handover notes for Kimi, operators, and future agents."),
            ("WP-Z-03", "Stewardship Runbook", "Document post-Atlas care routines, review cadence, maintenance windows, and authority boundaries."),
            ("WP-Z-04", "Archive Retention Package", "Package retained evidence, final documents, commit hashes, and audit indexes for 90-day retention."),
            ("WP-Z-05", "V5 to V6 Transition Brief", "Prepare the next-horizon brief with open risks, deferred work, and candidate V6 trains."),
        ],
    },
}

AC_TEMPLATE = [
    "Implementation is present in the repository with code, route, component, script, or document changes directly traceable to this WP.",
    "The user-facing or operator-facing workflow can be exercised end to end without mocked success states.",
    "RBAC and protected-route behavior are enforced where the WP touches authenticated, admin, or operator surfaces.",
    "Input validation, empty state, loading state, and failure state behavior are implemented where relevant.",
    "Evidence is collected under /evidence with the correct EP/date/category naming convention and a git commit hash reference.",
    "Automated test or scripted verification exists and records PASS/FAIL output for the WP.",
    "Build verification passes with bun run build, and lint verification passes with bun run lint before submission.",
    "The WP avoids reopening or weakening earlier accepted trains, routes, schemas, or evidence commitments.",
    "Kimi can independently review the result using the linked evidence, route, command output, and acceptance notes.",
    "Closure language states the final PASS/FAIL status, residual risks, and next authorized step.",
]

STATUS_COLOR = {
    "CLOSED": RGBColor(0x16, 0x7A, 0x3A),
    "ACTIVE": RGBColor(0x00, 0x6B, 0xFF),
    "PENDING": RGBColor(0x88, 0x88, 0x88),
}


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def body(doc, text):
    return doc.add_paragraph(text)


def add_table_row(table, cells, bold_first=False, header=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        run = cell.paragraphs[0].runs
        if run:
            run[0].bold = header or (bold_first and i == 0)
    return row


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_colored_text(para, text, color: RGBColor, bold=True):
    run = para.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    return run


def build():
    doc = Document()

    # Title
    t = doc.add_heading("ONX ATLAS V5 COMPLETE EXECUTION PACKAGE", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Trains L through Z  |  15 Trains  |  75 Work Packages  |  750 Acceptance Criteria")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph("")

    # Metadata table
    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    pairs = [
        ("Document", "ONX_ATLAS_V5_COMPLETE_EXECUTION_PACKAGE_v1.1.docx"),
        ("Version", "v1.1"),
        ("Date", "2026-06-24"),
        ("Status", "ACTIVE — Train N Current"),
        ("Executor", "Husam Yildirim"),
        ("Authority", "Kimi Review"),
    ]
    for row, (k, v) in zip(meta.rows, pairs):
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph("")

    # Part 1
    heading(doc, "Part 1: Atlas V5 Completion Overview", 1)
    body(doc, "This package finalizes the remaining Atlas V5 execution sequence. Trains L through M are CLOSED. Train N is the current active train. Trains O through Z are pending.")

    doc.add_paragraph("")
    scope = doc.add_table(rows=5, cols=3)
    scope.style = "Table Grid"
    headers = ["Scope Item", "Count", "Rule"]
    for i, h in enumerate(headers):
        scope.rows[0].cells[i].text = h
        if scope.rows[0].cells[i].paragraphs[0].runs:
            scope.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    rows_data = [
        ("Trains", "15", "Each train passes only at 50/50 criteria."),
        ("Work Packages", "75", "Each WP passes only at 10/10 criteria."),
        ("Acceptance Criteria", "750", "Binary PASS/FAIL. No partial credit."),
        ("Evidence", "Required for every WP", "Stored under /evidence/<EP-ID>/<YYYY-MM-DD>/<category>/."),
    ]
    for i, row_data in enumerate(rows_data, 1):
        for j, val in enumerate(row_data):
            scope.rows[i].cells[j].text = val

    doc.add_paragraph("")
    status_tbl = doc.add_table(rows=16, cols=3)
    status_tbl.style = "Table Grid"
    hdrs = ["Train", "Name", "Status"]
    for i, h in enumerate(hdrs):
        status_tbl.rows[0].cells[i].text = h
        if status_tbl.rows[0].cells[i].paragraphs[0].runs:
            status_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row_idx, (letter, (status, name)) in enumerate(TRAIN_STATUS.items(), 1):
        status_tbl.rows[row_idx].cells[0].text = f"Train {letter}"
        status_tbl.rows[row_idx].cells[1].text = name
        status_tbl.rows[row_idx].cells[2].text = status
        p = status_tbl.rows[row_idx].cells[2].paragraphs[0]
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = STATUS_COLOR[status]

    doc.add_page_break()

    # Parts 2-4: Each train
    for letter, (status, name) in TRAIN_STATUS.items():
        train_data = TRAINS[letter]
        heading(doc, f"Train {letter}: {name}", 1)

        # Status badge
        sp = doc.add_paragraph()
        add_colored_text(sp, f"STATUS: {status}", STATUS_COLOR[status])
        if status == "CLOSED":
            sp.add_run(" — Evidence committed. Kimi verified. Final verification document issued.")
        elif status == "ACTIVE":
            sp.add_run(" — Current train. Execute WP-N-01 through WP-N-05 in sequence.")

        body(doc, f"Master Objective: {train_data['objective']}")
        body(doc, f"Primary Routes / Surfaces: {train_data['routes']}")
        body(doc, f"Dependencies: {train_data['deps']}")
        body(doc, f"Acceptance Target: 50/50 PASS across WP-{letter}-01 through WP-{letter}-05")
        doc.add_paragraph("")

        for wp_id, wp_name, wp_desc in train_data["wps"]:
            heading(doc, f"{wp_id}: {wp_name}", 2)
            body(doc, wp_desc)
            doc.add_paragraph("")

            heading(doc, f"Acceptance Criteria ({wp_id}-AC-01 through {wp_id}-AC-10):", 3)
            for ac_idx, ac_text in enumerate(AC_TEMPLATE, 1):
                p = doc.add_paragraph(style="List Number")
                p.add_run(f"AC-{wp_id.split('-', 1)[1]}-{ac_idx:02d}: ").bold = True
                p.add_run(ac_text)

            doc.add_paragraph("")
            heading(doc, "Required Evidence:", 3)
            ev_prefix = f"EV-CODE_{wp_id}"
            ev_items = [
                f"EV-CODE_{wp_id}_<YYYYMMDD>_<brief>.*",
                f"EV-TEST_{wp_id}_<YYYYMMDD>_<brief>.*",
                f"EV-ACPT_{wp_id}_<YYYYMMDD>_<brief>.*",
                f"EV-DEPL_{wp_id}_<YYYYMMDD>_<brief>.* (when deployment or route behavior is touched)",
            ]
            for ev in ev_items:
                doc.add_paragraph(ev, style="List Bullet")

            doc.add_paragraph("")

        closure = doc.add_paragraph()
        closure.add_run(f"Train {letter} Closure Definition: ").bold = True
        closure.add_run(
            f"Train {letter} closes only when WP-{letter}-01 through WP-{letter}-05 each achieve 10/10 PASS, "
            f"evidence is committed, Kimi verifies completeness, and ONX_TRAIN_{letter}_FINAL_VERIFICATION_v1.0.docx is issued."
        )

        doc.add_page_break()

    # Part 5: Master Acceptance Matrix
    heading(doc, "Part 5: Master Acceptance Matrix", 1)
    body(doc, "All 15 trains (L through Z), each covering WP-01 to WP-05, 50 criteria each.")
    doc.add_paragraph("")

    mat = doc.add_table(rows=1 + 15, cols=5)
    mat.style = "Table Grid"
    mat_headers = ["Train", "WPs", "Criteria", "Status", "Reviewer"]
    for i, h in enumerate(mat_headers):
        mat.rows[0].cells[i].text = h
        if mat.rows[0].cells[i].paragraphs[0].runs:
            mat.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for row_idx, (letter, (status, name)) in enumerate(TRAIN_STATUS.items(), 1):
        mat.rows[row_idx].cells[0].text = f"Train {letter}"
        mat.rows[row_idx].cells[1].text = f"WP-{letter}-01 through WP-{letter}-05"
        mat.rows[row_idx].cells[2].text = "50"
        p = mat.rows[row_idx].cells[3].paragraphs[0]
        status_label = "CLOSED / PASSED" if status == "CLOSED" else ("ACTIVE — NEXT" if status == "ACTIVE" else "Pending execution")
        run = p.add_run(status_label)
        run.bold = True
        run.font.color.rgb = STATUS_COLOR[status]
        mat.rows[row_idx].cells[4].text = "Kimi"

    doc.add_paragraph("")

    # Part 6: Evidence Standard
    heading(doc, "Part 6: Master Evidence Standard", 1)
    ev_tbl = doc.add_table(rows=8, cols=3)
    ev_tbl.style = "Table Grid"
    ev_headers = ["Category", "Purpose", "Required Pattern"]
    for i, h in enumerate(ev_headers):
        ev_tbl.rows[0].cells[i].text = h
        if ev_tbl.rows[0].cells[i].paragraphs[0].runs:
            ev_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    ev_rows = [
        ("Code", "Implementation proof", "EV-CODE_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Test", "Automated or scripted verification", "EV-TEST_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Security", "Auth, RBAC, or safety proof", "EV-SEC_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Deployment", "Build, staging, release, or rollback proof", "EV-DEPL_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Runtime", "Health, latency, smoke, or monitoring proof", "EV-RUN_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Acceptance", "Checklist and reviewer decision", "EV-ACPT_<WP-ID>_<YYYYMMDD>_<brief>.*"),
        ("Closure", "Train or program closure order", "EV-CLSR_<EP-ID>_<YYYYMMDD>_<brief>.*"),
    ]
    for i, row_data in enumerate(ev_rows, 1):
        for j, val in enumerate(row_data):
            ev_tbl.rows[i].cells[j].text = val

    doc.add_paragraph("")

    # Part 7: Dependency Graph
    heading(doc, "Part 7: Master Dependency Graph", 1)
    body(doc, "K → L → M → N → O → P → Q → R → S → T → U → V → W → X → Y → Z")
    body(doc, "No train begins before predecessor closure. No work package begins before predecessor WP closure inside the same train.")
    doc.add_paragraph("")

    # Part 8: Execution Sequencing Rules
    heading(doc, "Part 8: Execution Sequencing Rules", 1)
    rules = [
        "Execute trains in strict alphabetical order from N through Z (Trains L and M are CLOSED).",
        "Execute WPs sequentially from WP-01 through WP-05 inside each train.",
        "Before WP submission, run lint, test, build, and route checks relevant to the WP.",
        "Every acceptance decision must cite concrete evidence and a commit hash.",
        "Any failed criterion blocks WP closure until remediation evidence is committed.",
    ]
    for rule in rules:
        doc.add_paragraph(rule, style="List Number")

    doc.add_paragraph("")

    # Part 9: Kimi Verification Instructions
    heading(doc, "Part 9: Kimi Verification Instructions", 1)
    kimi_rules = [
        "Verify every WP has exactly 10 acceptance criteria.",
        "Verify the evidence files exist, are current, and map to the claimed WP.",
        "Score each criterion as PASS or FAIL only.",
        "Reject verbal confirmation without a reviewable artifact.",
        "Escalate any ambiguity, missing evidence, or scope drift before closure.",
    ]
    for rule in kimi_rules:
        doc.add_paragraph(rule, style="List Number")

    doc.add_paragraph("")

    # Part 10: Finalization Clause
    heading(doc, "Part 10: Finalization Clause", 1)
    p = doc.add_paragraph()
    p.add_run("Executor: ").bold = True
    p.add_run("Husam Yildirim")
    p2 = doc.add_paragraph()
    p2.add_run("Authority: ").bold = True
    p2.add_run("Kimi Review")
    doc.add_paragraph("")
    body(
        doc,
        "Atlas V5 is not complete until all 15 trains, all 75 work packages, and all 750 criteria are scored PASS "
        "with evidence committed to the repository. The authorized completion target is 750/750 PASS. "
        "Trains L and M are CLOSED. Train N is the current active train. Proceed sequentially."
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
